import json

from flask import Flask, render_template, send_file
from flask import jsonify, request
from flask import make_response
import pymysql
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx import Document
from io import BytesIO
from flask_cors import CORS

from datetime import datetime

app = Flask(__name__)
CORS(app)
'''
플라스크 설치 방법 !! 
https://krksap.tistory.com/1750
'''
app.config['JSON_AS_ASCII'] = False;
name = "";
birth = "";
startDate = "";
endDate = "";

@app.route('/')
def home():
    return render_template('index.html');

# api 테스트
@app.route("/sign", methods=['post'])
def sign():
    user = request.json;
    response = {
        'name' : user['name'],
        'email' : user['email'],
        'password' : user['password']
    }
    result = json.dumps(response, ensure_ascii=False);
    res = make_response(result);
    return res, 200

@app.route("/classEnd", methods=['POST'])
def classEnd():
    student = request.json;
    response = {
        'name': student['name'],
        'birth': student['birth'],
        'startDate': student['startDate'],
        'endDate': student['endDate']
    };
    #print(response);
    global name
    global birth
    global startDate
    global endDate
    name = response['name'];
    birth = response['birth'];
    startDate = response['startDate']
    endDate = response['endDate']
    res = {
        "result": "ok",
    };
    result = jsonify(res)
    return result, 200

@app.route("/sus")
def success():
    content = "위 학생은 코리아IT아카데미 파이썬 수업을 매우 성실히 임하였기에 이 증서를 부여함"
    # 워드 문서 생성
    document = Document()
    '''p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('수료증')
    r.font.size = Pt(25)
    r.bold = True'''

    p = document.add_heading('Python 수료증', level=0);
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER;

    n = document.add_paragraph();
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    n1 = n.add_run("이름 : ");
    n1.font.size = Pt(15);
    n2 = n.add_run(name);
    n2.font.size = Pt(15);

    b = document.add_paragraph();
    b.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    b1 = b.add_run("생년월일 : ");
    b1.font.size = Pt(15);
    b2 = b.add_run(birth);
    b2.font.size = Pt(15);

    s = document.add_paragraph();
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s1 = s.add_run("개강일자 : ");
    s1.font.size = Pt(15);
    b2 = s.add_run(startDate);
    b2.font.size = Pt(15);

    e = document.add_paragraph();
    e.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    e1 = e.add_run("종강일자 : ");
    e1.font.size = Pt(15);
    e2 = e.add_run(endDate);
    e2.font.size = Pt(15);
    e.add_run("\n\n\n\n\n");

    c = document.add_paragraph();
    c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = c.add_run(content);
    r1.font.size = Pt(20);
    c.add_run("\n\n\n\n\n\n");

    # 현재 날짜와 시간 얻기
    now = datetime.now();
    current_year = now.year;
    current_month = now.month;
    current_day = now.day;
    day = document.add_paragraph();
    day.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    day_y = day.add_run(str(current_year));
    day_y.font.size = Pt(18);
    day.add_run("년   ")
    day_m = day.add_run(str(current_month));
    day_m.font.size = Pt(18);
    day.add_run("월   ")
    day_d = day.add_run(str(current_day));
    day_d.font.size = Pt(18);
    day.add_run("일");
    day.add_run("\n\n\n\n");

    end = document.add_paragraph();
    end.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    end1 = end.add_run("코리아IT아카데미 윤준형 강사");
    end1.font.size = Pt(20);
    end1.bold = True;

    # BytesIO를 사용하여 메모리에 파일을 저장
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Flask response로 파일을 반환
    return send_file(file_stream, download_name="generated_document.docx", as_attachment=True)



if __name__ == '__main__':
    app.run(port=4000);